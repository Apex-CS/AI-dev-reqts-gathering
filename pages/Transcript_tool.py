import io
from typing import Dict

import streamlit as st
from dotenv import load_dotenv
from langchain_core.chat_history import InMemoryChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
import docx
from fpdf import FPDF
import markdown
from bs4 import BeautifulSoup
from docx import Document
import json
import src.functions.helpers as helpers
import PyPDF2

import src.classes.prompt_templates as pt
import src.functions.ado_connector as ado_connector
import src.functions.jira_connector as jira_connector

from src.functions import utility_functions
import re
import datetime

from src.functions.settings import (
    save_project_details,
    save_rqm_tool_details,
    get_all_rqm_tool_details,
    get_projects_details,
    save_remove_work_items_project
)

# --- Constants and Config ---
load_dotenv()
DEFAULT_SESSION_STATE = {
    "messages": [],
    "chat_history_store": {},
    "summary": None,
    "user_stories": None,
    "filename": "",
    "connection_ado_default": {
        "base_url": "",
        "tool_type": "",
        "tool_name": "",
        "user_email": "",
        "project_name": "",
        "personal_access_token": ""
    },
}
DOCX_TYPES = [
    "application/docx",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
]

# --- Session State Initialization ---
for key, value in DEFAULT_SESSION_STATE.items():
    st.session_state.setdefault(key, value)

# --- Utility Functions ---
def safe_text(text):
    return text.encode('latin-1', 'replace').decode('latin-1')

def extract_text_from_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file_bytes):
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = []
    for page in reader.pages:
        text.append(page.extract_text())
    return "\n".join(text)

def render_messages(messages):
    for message in messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

def add_message(role, content):
    st.session_state.messages.append({"role": role, "content": content})

def export_chat_to_pdf(messages):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for message in messages:
        role = message["role"].capitalize()
        html = markdown.markdown(message["content"])
        soup = BeautifulSoup(html, "html.parser")
        content = soup.get_text()
        pdf.multi_cell(0, 10, safe_text(f"{role}:\n{content}\n"), align="L")
        pdf.ln(2)
    return pdf.output(dest='S').encode('latin-1', 'replace')

# --- Streamlit App Logic ---
def render():
    
    connector = None
    st.title("Requirements analysis")
    if st.session_state.doc_added:
        st.toast("Document added. You can now ask questions related to the document.")
        handle_user_input()
    else:
        st.warning("Please upload a document to get started.")
        handle_file_upload()

def handle_file_upload():
    uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx"])
    documents = ""
    if uploaded_file:
        st.session_state.filename = uploaded_file.name.rsplit('.', 1)[0]
        if uploaded_file.type == "application/pdf":
            documents = extract_text_from_pdf(uploaded_file.read())
        elif uploaded_file.type in DOCX_TYPES:
            documents = extract_text_from_docx(uploaded_file.read())
    submit = st.button("Submit", type="primary")
    if submit and uploaded_file:
        prompt_text = pt.initial_transcription_template.format(documents=documents)
        response = helpers.invoke_with_history(
            prompt_text,
            session_id="transcript_session"
        )
        try:
            response_json = json.loads(response[response.index('{'):response.rindex('}')+1])
        except Exception as e:
            print("Failed to parse response as JSON:", e)
            response_json = {}

        st.session_state.summary = response_json.get("summary", None)
        st.session_state.user_stories = response_json.get("user_stories", None)
        
        #add_message("user", prompt_text)
        add_message("assistant", st.session_state.summary)
        for story in st.session_state.user_stories or []:
            formated_story = f"**Title:** {story.get('title', '')}\n\n**Description:** {story.get('description', '')}\n\n**Acceptance Criteria:**\n" + "\n".join([f"- {ac}" for ac in story.get('acceptance_criteria', [])])
            add_message("assistant", formated_story)
        handle_user_input()
        
        

def handle_user_input():
    render_messages(st.session_state.messages)
    user_input = st.chat_input(
        "Type your message here...",
        accept_file=True,
        file_type=["pdf", "docx"],
    )
    documents = ""
    if user_input:
        question = getattr(user_input, "text", user_input)
        add_message("user", question)
        prompt_text = pt.answer_template.format(question=question, documents=documents)
        response = helpers.invoke_with_history(
            prompt_text,
            session_id="transcript_session"
        )
        add_message("assistant", response)
        render_messages(st.session_state.messages)
    if st.button("📄 Export Chat to PDF", type="primary"):
        pdf_output = export_chat_to_pdf(st.session_state.messages)
        st.download_button(
            label="Download Chat as PDF",
            data=pdf_output,
            file_name="chat_history.pdf",
            mime="application/pdf"
        )
    
    st.toast("Document added. You can now ask questions related to the document.")
    st.session_state.doc_added = True
    if st.session_state.summary:
        doc = Document()
        doc.add_heading("Document Summary", level=1)
        # Convert markup (markdown) to HTML, then parse and add to docx
        html = markdown.markdown(st.session_state.summary)
        soup = BeautifulSoup(html, "html.parser")
        for element in soup.children:
            if element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "ul":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "ol":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Number")
            elif element.name and element.name.startswith("h") and element.name[1:].isdigit():
                level = int(element.name[1:])
                doc.add_heading(element.get_text(), level=level)
            else:
                doc.add_paragraph(element.get_text())
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        st.download_button(
        label="Download Summary as Word Document",
        data=doc_bytes, 
        type="primary",
        file_name="document_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    project_names = []
    toolsNames = []
    for project in get_projects_details(utility_functions.SETTINGS_DB):
        project_names.append(project['project_name'])
    if st.session_state.user_stories:
        st.write("## Save Items in ALM Tool")
        selected_project = st.selectbox("Select Project", project_names)
        if selected_project:
            print(f"### Selected project: {selected_project}")

            alm_tools = get_all_rqm_tool_details(
                utility_functions.SETTINGS_DB, selected_project
            )
            toolsNames = []
            for tool in alm_tools:
                toolsNames.append(tool['tool_name'])

            st.write("### Select ALM Tool")
        selected_tool = st.selectbox("Select ALM Tool", toolsNames, key="alm_tool_select")
        if selected_tool:
            tool = next((t for t in alm_tools if t['tool_name'] == selected_tool), None)
            if tool['tool_type'] == "Jira":
                connector = jira_connector.JiraConnector()
            elif tool['tool_type'] == "ADO":
                connector = ado_connector.AdoConnector()
            connector.change_connection(
                tool["url"],
                tool["tool_name"],
                tool["pat"],
                tool["user_email"]
            )
            
            st.info(f"Connection details updated for: {selected_tool}")

            if st.button("Confirm and Save Work Items", type="primary"):
                st.toast("Saving work items to Azure DevOps...")
                for story in st.session_state.user_stories:
                    title = story.get("title", "No Title")
                    description = story.get("description", "")
                    acceptance_criteria = "\n".join([f"- {ac}" for ac in story.get("acceptance_criteria", [])])
                    connector.add_work_item(title, description, acceptance_criteria, tool["tool_name"])
                st.success("Work items have been saved to Azure DevOps.")
                
            if st.button("💾 Save as Project", type="primary"):
                st.toast("Saving as a project...")
                # Extract context section from summary
                context_text = ""
                if st.session_state.summary:
                    match = re.search(r'#### Context:(.*?)####', st.session_state.summary, re.DOTALL)
                    if match:
                        context_text = match.group(1).strip()
                project_name = f"{datetime.datetime.now().strftime('%Y-%m-%d')} - {st.session_state.filename}"
                project_description = context_text
                save_project_details(
                    db_path=utility_functions.SETTINGS_DB,
                    project_name=project_name,
                    project_description=project_description,
                    project_summary=st.session_state.summary
                )
                conn_info = st.session_state.connection_ado_default
                save_rqm_tool_details(
                    db_path=utility_functions.SETTINGS_DB,
                    project=project_name,
                    tool_type=tool["tool_type"],
                    rqm_type=tool["rqm_type"],
                    url=tool["url"],
                    tool_name=tool["tool_name"],
                    pat=tool["pat"],
                    user_email=tool["user_email"]
                )
                for story in st.session_state.user_stories:
                    title = story.get("title", "No Title")
                    description = story.get("description", "")
                    acceptance_criteria = "\n".join([f"- {ac}" for ac in story.get("acceptance_criteria", [])])
                    last_item = connector.add_work_item(title, description, acceptance_criteria, tool["tool_name"])
                    save_remove_work_items_project(
                        db_path=utility_functions.SETTINGS_DB,
                        project_name=project_name,
                        rqm_name=tool["tool_name"],
                        work_item_id=last_item.id
                    )

                st.success("Work items have been saved to Azure DevOps.")

                st.session_state["load_tree"] = True
                st.rerun()
                st.success(f"Project '{project_name}' added successfully.")